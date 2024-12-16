import requests
import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

response = requests.get('https://edition.cnn.com/travel/article/scenic-airport-landings-2020/index.html')
page = BeautifulSoup(response.content,'html.parser')

title = page.find('h1',{'id':'maincontent'}).text.strip()

byline_name = page.find('span',{'class':'byline__name'}).text

updated = page.find('div',{'class':'timestamp vossi-timestamp'}).text.strip().replace("\n    ","")

paragraphs = page.find_all('p',{'class':'paragraph'})
paragraphs_list = []
for p in paragraphs:
    paragraphs_list.append(p.text.strip())

head_of_content = page.find('cite',{'class':'source__cite'}).find('span',{'class':'source__text'}).text + " - " + paragraphs_list[0]

content = paragraphs_list[1:-10]

head_of_top_10 = page.find('h2',{'class':'subheader'}).text.strip()


gallery = page.find('div',{'class':'gallery-inline__slides'})
all_data = gallery.find_all("div",{'class':'image image__hide-placeholder'})

data_list_dict = []
for i in all_data:
    image = i.find('div',{'class':'image__container'}).find('picture',{'class':'image__picture'}).find('img').attrs['src']
    name = i.find('div',{'class':'image__metadata'}).find('div',{'class':'image__caption attribution'}).find('span').text.split(":")[0]
    cap = i.find('div',{'class':'image__metadata'}).find('div',{'class':'image__caption attribution'}).find('span').text.split(":")[1]

    if name == "World's most scenic airports":
        name = "10. Nadi International Airport, Fiji"

    data_dict = {
        'Name':name,
        'Caption':cap,
        'Image':image
    }

    data_list_dict.append(data_dict)


os.makedirs("scraped_images", exist_ok=True)
for index,value in enumerate(data_list_dict):
    response = requests.get(value['Image'])

    filename = f"image_{index+1}.jpg"

    file_path = os.path.join("scraped_images", filename)

    with open(file_path, "wb") as f:
        f.write(response.content)

    print("Downloaded:", file_path)


doc = Document()
head = doc.add_heading(title,level=1)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("by"+" "+byline_name,level=3)
doc.add_heading(updated,level=3)
doc.add_paragraph('')
doc.add_heading(head_of_content,level=2)
for i in content:
    doc.add_paragraph(i)
doc.add_paragraph('')
doc.add_heading(head_of_top_10,level=1)
doc.add_paragraph('')
for index,value in enumerate(data_list_dict):
    doc.add_picture(f'scraped_images\\image_{index+1}.jpg', width=Inches(5), height=Inches(3))
    doc.add_heading(value['Name'][3:],level=3)
    doc.add_paragraph(value['Caption'])
doc.save("scraped_document.docx")